from flask import Flask, request, send_from_directory
from dotenv import load_dotenv
import openai
from docx import Document
from docx.shared import Inches
import json
import os

app = Flask(__name__)

# Set up OpenAI API credentials
load_dotenv()
openai.api_key = os.environ.get('OAI_SECRET_KEY')

# Define a route for the home page
# @app.route('/')
# def home():
#     return render_template('home.html')



@app.route('/process', methods=['POST'])
def process():
    # Get user input text
    user_text = request.form['user_text']

    # Add predefined text
    predefined_text = """  
    Using the text before this sentence, I want you to generate a marketing brief, seperated into these topics: Background, Objective, Target Audience, and Brand Guidelines. Separate each topic with an empty line and do not include empty lines anywhere else. Be as thorough as possible and include as much information as possible. Make sure all responses are complete sentences and include atleast 4 sentences for each topic.
      """
    combined_text = predefined_text + user_text

    # Send combined text to OpenAI GPT-3 API
    response = openai.Completion.create(
        engine="text-davinci-003",
        prompt=combined_text,
        temperature=0.9,
        max_tokens=2048,
        n = 1,
        stop=None
    )
    output_text = response.choices[0].text

    # print(output_text)

    # Extract relevant sections from output
    background_text = ""
    objective_text = ""
    target_audience_text = ""
    brand_guidelines_text=""
    for line in output_text.split("\n\n"):
        if "background:" in line.lower():
            background_text += line + "\n"
        elif "objective:" in line.lower():
            objective_text += line + "\n"
        elif "target audience:" in line.lower():
            target_audience_text += line + "\n"
        elif "brand guidelines:" in line.lower():
            brand_guidelines_text += line + "\n"

    # Render output template
    return json.dumps({
        'background': background_text, 
        'objective': objective_text, 
        'target_audience': target_audience_text,
        'brand_guidelines': brand_guidelines_text}
    )

@app.route('/download', methods=[ 'POST' ])
def download():
    background_text = request.form['background']
    objective_text = request.form['objective']
    target_audience_text = request.form['target_audience']
    brand_guidelines_text = request.form['brand_guidelines']

    document = Document()

    style = document.styles['Normal']
    style.font.name='Calibri'

    # Add header and attach image to it
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture('./assets/OmniLogo.png', width=Inches(1))

    document.add_heading('AI-Creative Assistant Brief', 0)
    document.add_picture('./assets/briefoLogo.png', width=Inches(2))
    
    document.add_heading('Background:', level=1)
    document.add_paragraph(background_text)

    document.add_heading('Objective:', level=1)
    document.add_paragraph(objective_text)
    
    document.add_heading('Target Audience:', level=1)
    document.add_paragraph(target_audience_text)
    
    document.add_heading('Brand Guidelines:', level=1)
    document.add_paragraph(brand_guidelines_text)

    document.save('output.docx')

    try:
        return send_from_directory('./', 'output.docx', as_attachment=True)
    except Exception as e:
        return str(e)

if __name__ == '__main__':
    app.run(debug=True, port=5000)